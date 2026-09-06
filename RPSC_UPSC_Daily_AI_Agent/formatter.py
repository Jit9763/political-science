import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from master_library import MasterNotesLibrary

class NotesFormatter:
    def __init__(self, output_dir=None):
        if output_dir is None:
            output_dir = os.path.join(os.path.dirname(__file__), "Output_Notes")
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.master_kb = MasterNotesLibrary()

    def render_master_syllabus_html(self):
        """Render complete interactive RAS Master Syllabus Wiki HTML website."""
        topics_data = self.master_kb.get_all_topics()
        last_updated = self.master_kb.data.get("metadata", {}).get("last_updated", "Recently")

        today_str = datetime.now().strftime("%Y-%m-%d")
        recent_updates_html = ""
        recent_count = 0

        paper_sections_html = ""
        for paper_name, units in topics_data.items():
            units_html = ""
            for unit_name, items in units.items():
                items_html = ""
                if not items:
                    items_html = '<div class="empty-topic">नवीनतम दैनिक समाचारों से इस टॉपिक का अध्ययन जल्द ही अद्यतन होगा।</div>'
                else:
                    for item in items:
                        updates_list = ""
                        for u in item.get("updates_history", []):
                            u_notes = u.get('update_notes', '').replace('\n', '<br>')
                            updates_list += f"""
                            <div class="update-entry">
                                <span class="update-date">📅 {u.get('date', '')}</span>
                                <div class="update-text">{u_notes}</div>
                            </div>
                            """
                        status_text = item.get('current_status', '').replace('\n', '<br>')
                        
                        # Collect recent updates for top banner
                        if item.get('last_updated') == today_str or (recent_count < 3 and item.get('last_updated')):
                            recent_count += 1
                            recent_updates_html += f"""
                            <div class="syllabus-topic-card" style="border-left:6px solid #dc2626; background:#fff1f2;">
                                <div class="topic-header">
                                    <span class="topic-name">🔥 [{paper_name}] {item.get('title', '')}</span>
                                    <span class="badge" style="background:#dc2626; color:white; font-weight:bold;">ताज़ा अपडेट: {item.get('last_updated', '')}</span>
                                </div>
                                <div class="current-status-box">
                                    <strong>अद्यतन स्थिति एवं मुख्य तथ्य:</strong>
                                    <p>{status_text}</p>
                                </div>
                            </div>
                            """

                        items_html += f"""
                        <div class="syllabus-topic-card">
                            <div class="topic-header">
                                <span class="topic-name">📌 {item.get('title', '')}</span>
                                <span class="badge badge-update">अद्यतन तिथि: {item.get('last_updated', '')}</span>
                            </div>
                            <div class="current-status-box">
                                <strong>वर्तमान स्थिति एवं अद्यतन तथ्य:</strong>
                                <p>{status_text}</p>
                            </div>
                            <details class="history-details">
                                <summary>📜 अपडेट इतिहास एवं पूर्व विवरण देखें</summary>
                                <div class="history-body">{updates_list}</div>
                            </details>
                        </div>
                        """

                units_html += f"""
                <div class="unit-block">
                    <h3 class="unit-header">📘 {unit_name}</h3>
                    {items_html}
                </div>
                """

            paper_sections_html += f"""
            <div id="{paper_name.replace(' ', '_')}" class="paper-tab-content">
                <h2 class="paper-title">🏛️ RAS {paper_name} - विस्तृत पाठ्यक्रम एवं मास्टर नोट्स</h2>
                {units_html}
            </div>
            """

        wiki_html = f"""<!DOCTYPE html>
<html lang="hi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>RAS Master Syllabus Wiki - RPSC 4 Papers Knowledge Base</title>
    <link href="https://fonts.googleapis.com/css2?family=Noto+Sans+Devanagari:wght@400;500;600;700;800;900&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #0f172a;
            --primary-light: #2563eb;
            --accent-green: #059669;
            --accent-purple: #7c3aed;
            --bg-main: #f8fafc;
            --card-bg: #ffffff;
            --text-dark: #020617;
            --border-color: #cbd5e1;
        }}
        body {{
            font-family: 'Noto Sans Devanagari', sans-serif;
            background-color: var(--bg-main);
            color: var(--text-dark);
            margin: 0;
            padding: 0;
            line-height: 1.8;
            font-size: 18px;
        }}
        .top-navbar {{
            background: linear-gradient(135deg, #0f172a 0%, #1e3a8a 100%);
            color: white;
            padding: 25px 40px;
            box-shadow: 0 10px 20px rgba(0,0,0,0.15);
        }}
        .top-navbar h1 {{
            margin: 0 0 8px 0;
            font-size: 30px;
            font-weight: 900;
        }}
        .top-navbar p {{
            margin: 0;
            opacity: 0.9;
            font-size: 17px;
        }}
        .container {{
            max-width: 1250px;
            margin: 30px auto;
            padding: 0 20px;
        }}
        .tab-buttons {{
            display: flex;
            gap: 15px;
            margin-bottom: 30px;
            border-bottom: 3px solid #e2e8f0;
            padding-bottom: 10px;
        }}
        .tab-btn {{
            background: #ffffff;
            color: #334155;
            border: 2px solid #cbd5e1;
            padding: 12px 24px;
            border-radius: 12px;
            font-size: 19px;
            font-weight: 800;
            cursor: pointer;
            font-family: inherit;
            transition: all 0.2s;
        }}
        .tab-btn:hover, .tab-btn.active {{
            background: #2563eb;
            color: white;
            border-color: #2563eb;
            box-shadow: 0 4px 12px rgba(37,99,235,0.3);
        }}
        .paper-title {{
            font-size: 26px;
            font-weight: 900;
            color: #1e3a8a;
            margin-bottom: 25px;
            padding-bottom: 10px;
            border-bottom: 3px solid #3b82f6;
        }}
        .unit-block {{
            background: white;
            border-radius: 16px;
            padding: 25px;
            margin-bottom: 30px;
            box-shadow: 0 4px 10px rgba(0,0,0,0.05);
            border: 1px solid var(--border-color);
        }}
        .unit-header {{
            font-size: 22px;
            font-weight: 800;
            color: #0f766e;
            margin-top: 0;
            margin-bottom: 20px;
            padding-bottom: 8px;
            border-bottom: 2px solid #ccfbf1;
        }}
        .syllabus-topic-card {{
            background: #f8fafc;
            border-radius: 12px;
            padding: 20px;
            margin-bottom: 20px;
            border-left: 6px solid #2563eb;
            border: 1px solid #e2e8f0;
            border-left-width: 6px;
        }}
        .topic-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 12px;
        }}
        .topic-name {{ font-size: 20px; font-weight: 800; color: #0f172a; }}
        .badge-update {{ background: #059669; color: white; padding: 4px 12px; border-radius: 6px; font-size: 14px; font-weight: 700; }}
        
        .current-status-box {{
            background: #ffffff;
            padding: 16px;
            border-radius: 10px;
            border: 1px solid #cbd5e1;
            font-size: 18px;
        }}
        .history-details {{ margin-top: 12px; font-size: 16.5px; color: #475569; }}
        .history-body {{ margin-top: 10px; padding: 12px; background: #f1f5f9; border-radius: 8px; }}
        .update-entry {{ margin-bottom: 10px; padding-bottom: 8px; border-bottom: 1px dashed #cbd5e1; }}
        .update-date {{ font-weight: 700; color: #1e3a8a; font-size: 15px; }}
        .empty-topic {{ color: #94a3b8; font-style: italic; padding: 10px 0; }}

        .search-box {{
            width: 100%;
            padding: 14px 20px;
            font-size: 18px;
            border-radius: 12px;
            border: 2px solid #cbd5e1;
            margin-bottom: 25px;
            font-family: inherit;
            box-sizing: border-box;
        }}
    </style>
</head>
<body>
    <div class="top-navbar">
        <h1>📚 RPSC RAS 4 पेपर मास्टर सिलेबस नॉलेज बेस (Syllabus Wiki)</h1>
        <p>पाठ्यक्रम अनुसार व्यवस्थित अध्ययन सामग्री एवं स्वतः अद्यतन तथ्य | अंतिम अद्यतन: {last_updated}</p>
    </div>

    <div class="container">
        <input type="text" class="search-box" id="syllabusSearch" onkeyup="filterTopics()" placeholder="🔍 RAS सिलेबस टॉपिक, आयोग, नियम या विषय खोजें...">

        {f'''
        <div style="background:#fff1f2; border:2px solid #fecdd3; border-radius:14px; padding:20px; margin-bottom:25px;">
            <h2 style="margin:0 0 15px 0; color:#991b1b; font-size:22px; font-weight:900;">🔥 हाल ही में अद्यतन किए गए नवीन तथ्य (Recent Live Updates)</h2>
            {recent_updates_html}
        </div>
        ''' if recent_updates_html else ''}

        <div class="tab-buttons">
            <button class="tab-btn active" onclick="showPaper('Paper_1', this)">📘 Paper 1</button>
            <button class="tab-btn" onclick="showPaper('Paper_2', this)">📗 Paper 2</button>
            <button class="tab-btn" onclick="showPaper('Paper_3', this)">📙 Paper 3</button>
            <button class="tab-btn" onclick="showPaper('Paper_4', this)">📕 Paper 4</button>
        </div>

        {paper_sections_html}
    </div>

    <script>
        function showPaper(paperId, btn) {{
            let contents = document.querySelectorAll('.paper-tab-content');
            contents.forEach(c => c.style.display = 'none');
            document.getElementById(paperId).style.display = 'block';

            let buttons = document.querySelectorAll('.tab-btn');
            buttons.forEach(b => b.classList.remove('active'));
            btn.classList.add('active');
        }}

        function filterTopics() {{
            let input = document.getElementById('syllabusSearch').value.toLowerCase();
            let cards = document.querySelectorAll('.syllabus-topic-card');
            cards.forEach(card => {{
                let text = card.innerText.toLowerCase();
                card.style.display = text.includes(input) ? 'block' : 'none';
            }});
        }}

        // Default open Paper 1
        showPaper('Paper_1', document.querySelector('.tab-btn'));
    </script>
</body>
</html>
"""
        master_wiki_path = os.path.join(self.output_dir, "Master_Syllabus_Wiki.html")
        with open(master_wiki_path, "w", encoding="utf-8") as f:
            f.write(wiki_html)
        print(f"Master Syllabus Wiki HTML saved to: {master_wiki_path}")
        self.render_master_wiki_markdown(topics_data, self.master_kb.data.get("metadata", {}))
        return master_wiki_path

    def render_master_wiki_markdown(self, topics, metadata):
        """Generates GitHub Markdown Wiki file for repository Wiki tab."""
        md_content = f"# 📚 RPSC RAS & UPSC Master Syllabus Wiki\n\n"
        md_content += f"*अंतिम अद्यतन: {metadata.get('last_updated', '')}*\n\n"
        md_content += "---" + "\n\n"

        for paper_name, units in topics.items():
            md_content += f"## 📄 {paper_name}\n\n"
            for unit_name, items in units.items():
                if not items:
                    continue
                md_content += f"### 📘 {unit_name}\n\n"
                for item in items:
                    md_content += f"#### 📌 {item.get('title', '')}\n"
                    md_content += f"- **अद्यतन तिथि**: `{item.get('last_updated', '')}`\n"
                    md_content += f"- **वर्तमान स्थिति एवं तथ्य**: {item.get('current_status', '')}\n\n"
                    if item.get("updates_history"):
                        md_content += "<details><summary>📜 अपडेट इतिहास</summary>\n\n"
                        for u in item.get("updates_history", []):
                            md_content += f"- **{u.get('date', '')}**: {u.get('update_notes', '')}\n"
                        md_content += "</details>\n\n"

        md_path = os.path.join(self.output_dir, "Master_Syllabus_Wiki.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"Master Syllabus Wiki Markdown saved to: {md_path}")
        return md_path

    def render_daily_html(self, analysis_data):
        date_str = analysis_data.get('date', 'Today')
        
        prelims_facts = analysis_data.get('prelims_facts', [])
        mains_qs = analysis_data.get('mains_questions', [])
        editorials = analysis_data.get('editorial_deep_dive', [])
        yt_analysis = analysis_data.get('youtube_teacher_analysis', [])
        sujas = analysis_data.get('rajasthan_sujas_special', [])

        prelims_cards_html = ""
        for item in prelims_facts:
            tag_class = "badge-raj" if item.get('rajasthan_special') else "badge-pre"
            prelims_cards_html += f"""
            <div class="fact-card">
                <div class="card-header">
                    <span class="badge {tag_class}">{item.get('exam_tag', 'RAS/UPSC Pre')}</span>
                    <span class="topic-title">{item.get('topic', '')}</span>
                </div>
                <p class="fact-text">{item.get('fact', '')}</p>
            </div>
            """

        mains_qs_html = ""
        for q in mains_qs:
            marks = q.get('marks', 5)
            badge_class = "badge-5m" if marks == 5 else "badge-10m"
            
            if marks == 5:
                model_ans = q.get('model_answer', '').replace('\n', '<br>')
                answer_body = f"""
                <div class="answer-box">
                    <strong class="ans-label">उत्तर ढांचा (~50 शब्द):</strong>
                    <div class="answer-content">{model_ans}</div>
                </div>
                """
            else:
                body_text = q.get('body', '').replace('\n', '<br>')
                answer_body = f"""
                <div class="answer-box">
                    <div class="ans-section"><strong class="ans-label">भूमिका (Introduction):</strong><br>{q.get('intro', '')}</div>
                    <div class="ans-section"><strong class="ans-label">मुख्य भाग (Body & Rajasthan Context):</strong><br>{body_text}</div>
                    <div class="ans-section"><strong class="ans-label">निष्कर्ष (Conclusion):</strong><br>{q.get('conclusion', '')}</div>
                </div>
                """

            mains_qs_html += f"""
            <div class="mains-card">
                <div class="mains-header">
                    <span class="badge {badge_class}">{marks} अंक ({q.get('paper', 'Paper')})</span>
                    <span class="subject-tag">{q.get('subject', '')}</span>
                </div>
                <h3 class="question-text">प्रश्न: {q.get('question', '')}</h3>
                {answer_body}
            </div>
            """

        yt_html = ""
        for yt in yt_analysis:
            yt_html += f"""
            <div class="yt-card">
                <div class="yt-header">
                    <span class="badge badge-yt">📺 कोचिंग शिक्षक विश्लेषण</span>
                    <span class="topic-title">{yt.get('topic', '')}</span>
                </div>
                <div class="yt-explanation">
                    <strong>शिक्षक व्याख्यान एवं ट्रिक्स:</strong>
                    <p>{yt.get('teacher_explanation', '')}</p>
                </div>
                <div class="key-args">
                    <strong>क्लास के मुख्य बिंदु (Key Lecture Points):</strong>
                    <ul>
                        {"".join([f"<li>{pt}</li>" for pt in yt.get('key_takeaways', [])])}
                    </ul>
                </div>
                <div class="exam-tip-box"><strong>💡 परीक्षा टिप (Exam Tip):</strong> {yt.get('exam_tip', '')}</div>
            </div>
            """

        sujas_html = ""
        for s in sujas:
            sujas_html += f"""
            <div class="sujas-card">
                <div class="sujas-header">
                    <span class="badge badge-sujas">राजस्थान सुजस विशेष</span>
                    <strong>{s.get('department', 'DIPR Rajasthan')}</strong>
                </div>
                <h3>{s.get('title', '')}</h3>
                <ul>
                    {"".join([f"<li>{pt}</li>" for pt in s.get('key_points', [])])}
                </ul>
                <div class="relevance-box"><strong>RPSC महत्व:</strong> {s.get('rpsc_relevance', '')}</div>
            </div>
            """

        editorial_html = ""
        for ed in editorials:
            editorial_html += f"""
            <div class="editorial-card">
                <div class="editorial-header">
                    <span class="badge badge-ed">{ed.get('source', 'Editorial')}</span>
                    <span class="syllabus-tag">{ed.get('syllabus_topic', '')}</span>
                </div>
                <h3>{ed.get('title', '')}</h3>
                <p class="context"><strong>समसामयिक संदर्भ:</strong> {ed.get('context', '')}</p>
                <div class="key-args">
                    <strong>सम्पादकीय का पूरा सार एवं तर्क (Deep Dive Takeaways):</strong>
                    <ul>
                        {"".join([f"<li>{arg}</li>" for arg in ed.get('key_arguments', [])])}
                    </ul>
                </div>
                <div class="way-forward"><strong>आगे की राह (Way Forward):</strong> {ed.get('way_forward', '')}</div>
            </div>
            """

        html_content = f"""<!DOCTYPE html>
<html lang="hi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>RAS & UPSC Daily Notes - {date_str}</title>
    <link href="https://fonts.googleapis.com/css2?family=Noto+Sans+Devanagari:wght@400;500;600;700;800;900&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #1e1b4b;
            --primary-light: #3b82f6;
            --secondary: #0f766e;
            --accent-gold: #b45309;
            --accent-purple: #6b21a8;
            --accent-red: #9f1239;
            --bg-main: #f1f5f9;
            --card-bg: #ffffff;
            --text-dark: #020617;
            --border-color: #cbd5e1;
        }}
        body {{
            font-family: 'Noto Sans Devanagari', sans-serif;
            background-color: var(--bg-main);
            color: var(--text-dark);
            margin: 0;
            padding: 25px;
            line-height: 1.8;
            font-size: 18px;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
        }}
        .header {{
            background: linear-gradient(135deg, #09090b 0%, #1e1b4b 50%, #0f766e 100%);
            color: white;
            padding: 35px 40px;
            border-radius: 20px;
            box-shadow: 0 15px 30px -5px rgba(0,0,0,0.25);
            margin-bottom: 35px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        .header h1 {{ margin: 0 0 10px 0; font-size: 32px; font-weight: 900; }}
        .header p {{ margin: 0; opacity: 0.95; font-size: 19px; }}
        .print-btn {{
            background: #ffffff;
            color: #1e1b4b;
            border: none;
            padding: 14px 24px;
            border-radius: 12px;
            font-weight: 800;
            font-size: 17px;
            cursor: pointer;
            font-family: inherit;
            box-shadow: 0 6px 12px rgba(0,0,0,0.15);
        }}
        .section-title {{
            font-size: 26px;
            font-weight: 800;
            margin: 40px 0 20px 0;
            padding: 12px 20px;
            background: #ffffff;
            border-radius: 12px;
            border-left: 8px solid var(--primary-light);
            box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
            color: #0f172a;
        }}
        .badge {{ padding: 6px 14px; border-radius: 8px; font-size: 15px; font-weight: 800; color: white; display: inline-block; }}
        .badge-pre {{ background-color: var(--accent-gold); }}
        .badge-raj {{ background-color: var(--accent-purple); }}
        .badge-5m {{ background-color: #0284c7; }}
        .badge-10m {{ background-color: #4338ca; }}
        .badge-sujas {{ background-color: #0d9488; }}
        .badge-ed {{ background-color: var(--accent-red); }}
        .badge-yt {{ background-color: #dc2626; }}
        
        .fact-card, .mains-card, .sujas-card, .editorial-card, .yt-card {{
            background: var(--card-bg);
            border-radius: 16px;
            padding: 28px;
            margin-bottom: 25px;
            box-shadow: 0 6px 12px -2px rgba(0,0,0,0.08);
            border: 1px solid var(--border-color);
        }}
        .fact-card {{ border-left: 8px solid var(--accent-gold); }}
        .mains-card {{ border-left: 8px solid #4338ca; }}
        .sujas-card {{ border-left: 8px solid var(--accent-purple); }}
        .editorial-card {{ border-left: 8px solid var(--accent-red); }}
        .yt-card {{ border-left: 8px solid #dc2626; background: #fff1f2; }}
        
        .card-header, .mains-header, .sujas-header, .editorial-header, .yt-header {{ display: flex; align-items: center; gap: 14px; margin-bottom: 14px; }}
        .topic-title {{ font-weight: 800; font-size: 20px; color: #0f172a; }}
        .fact-text {{ margin: 0; color: var(--text-dark); font-size: 19px; line-height: 1.8; }}
        .question-text {{ margin: 12px 0 16px 0; color: #1e1b4b; font-size: 22px; font-weight: 800; }}
        .answer-box {{ background: #f8fafc; padding: 22px; border-radius: 12px; font-size: 18px; border: 1px solid #e2e8f0; }}
        .ans-label {{ font-size: 19px; color: #1e3a8a; display: inline-block; margin-bottom: 6px; }}
        .ans-section {{ margin-bottom: 16px; line-height: 1.8; }}
        .relevance-box, .way-forward {{ background: #ecfdf5; color: #065f46; padding: 14px 18px; border-radius: 10px; margin-top: 16px; font-size: 17.5px; border-left: 5px solid #10b981; font-weight: 600; }}
        .exam-tip-box {{ background: #fefce8; color: #854d0e; padding: 14px 18px; border-radius: 10px; margin-top: 16px; font-size: 17.5px; border-left: 5px solid #eab308; font-weight: 600; }}
        ul {{ margin: 10px 0; padding-left: 25px; }}
        li {{ margin-bottom: 8px; font-size: 18px; }}

        @media print {{
            .print-btn {{ display: none; }}
            body {{ background: white; padding: 0; font-size: 14pt; }}
            .container {{ max-width: 100%; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <div>
                <h1>RPSC RAS & UPSC दैनिक समसामयिकी एवं सम्पादकीय विश्लेषण</h1>
                <p>दिनांक: {date_str} | RAS Pre & Mains (5M & 10M) विशेष मास्टर नोट्स</p>
            </div>
            <div>
                <button class="print-btn" onclick="window.print()">🖨️ PDF प्रिंट करें</button>
            </div>
        </div>

        <div class="section-title">🟡 प्रारंभिक परीक्षा तथ्य (RAS & UPSC Prelims Tracker)</div>
        {prelims_cards_html if prelims_cards_html else '<p>आज के लिए प्रिलिम्स तथ्य प्रस्तुत हैं।</p>'}

        <div class="section-title">🔵 मुख्य परीक्षा उत्तर लेखन मॉडल सेट (RAS Mains - 5M & 10M)</div>
        {mains_qs_html if mains_qs_html else '<p>आज के लिए मुख्य परीक्षा प्रश्न मॉडल सेट।</p>'}

        {f'<div class="section-title">📺 कोचिंग शिक्षक विश्लेषण एवं यूट्यूब लाइव क्लास सार</div>{yt_html}' if yt_html else ''}

        <div class="section-title">🟣 राजस्थान सुजस एवं DIPR विशेष (RPSC Special)</div>
        {sujas_html if sujas_html else '<p>सुजस एवं राजस्थान राज्य सरकार की अद्यतन घोषणाएं।</p>'}

        <div class="section-title">🔴 सम्पादकीय विस्तृत विश्लेषण (The Hindu, Indian Express, ET & DownToEarth)</div>
        {editorial_html if editorial_html else '<p>दैनिक सम्पादकीय विस्तृत विश्लेषण प्रस्तुत है।</p>'}
    </div>
</body>
</html>
"""
        filepath = os.path.join(self.output_dir, f"RAS_UPSC_Notes_{date_str}.html")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html_content)
        
        # Render updated Master Syllabus Wiki HTML as well!
        self.render_master_syllabus_html()

        print(f"HTML Daily Notes saved to: {filepath}")
        return filepath

    def render_daily_docx(self, analysis_data):
        date_str = analysis_data.get('date', 'Today')
        doc = Document()
        
        heading = doc.add_heading(f'RAS & UPSC Daily Notes - {date_str}', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('1. प्रारंभिक परीक्षा तथ्य (Prelims Tracker)', level=1)
        for item in analysis_data.get('prelims_facts', []):
            p = doc.add_paragraph()
            p.add_run(f"• [{item.get('topic', '')}] ").bold = True
            p.add_run(item.get('fact', ''))

        doc.add_heading('2. मुख्य परीक्षा मॉडल उत्तर सेट (RAS Mains)', level=1)
        for q in analysis_data.get('mains_questions', []):
            p = doc.add_paragraph()
            p.add_run(f"प्रश्न [{q.get('marks', 5)} अंक - {q.get('paper', '')}]: {q.get('question', '')}\n").bold = True
            if q.get('marks') == 5:
                p.add_run(f"उत्तर: {q.get('model_answer', '')}\n")
            else:
                p.add_run(f"भूमिका: {q.get('intro', '')}\n")
                p.add_run(f"मुख्य भाग: {q.get('body', '')}\n")
                p.add_run(f"निष्कर्ष: {q.get('conclusion', '')}\n")

        docx_path = os.path.join(self.output_dir, f"RAS_UPSC_Notes_{date_str}.docx")
        doc.save(docx_path)
        print(f"Word DOCX Notes saved to: {docx_path}")
        return docx_path

if __name__ == '__main__':
    formatter = NotesFormatter()
    formatter.render_master_syllabus_html()
    print("Master Syllabus Wiki HTML rendering tested.")
